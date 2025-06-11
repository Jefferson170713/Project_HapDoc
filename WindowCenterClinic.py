import pandas as pd
from PyQt5.QtWidgets import QVBoxLayout
from PyQt5.QtWidgets import QLabel
from PyQt5.QtWidgets import QPushButton
from PyQt5.QtWidgets import QFrame
from PyQt5.QtWidgets import QFileDialog
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtWidgets import QHBoxLayout
from PyQt5.QtWidgets import QCheckBox
from PyQt5.QtWidgets import QButtonGroup
from PyQt5.QtWidgets import QLineEdit
from PyQt5.QtWidgets import QDialog
from PyQt5.QtWidgets import QTableWidget
from PyQt5.QtWidgets import QTableWidgetItem
from PyQt5.QtWidgets import QScrollArea
from PyQt5.QtWidgets import QWidget
from datetime import datetime
import os
from ARQUIVOS.Oracle_Jdbc.jdbc_permission import JdbcPermission 
from docx import Document # type: ignore
from docx.shared import Pt # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Para centralizar o texto # type: ignore
from docx.oxml import OxmlElement # type: ignore
from docx.oxml.ns import qn  # Para manipular bordas # type: ignore

class WindowCenterClinic:
    def __init__(self, parent=None):
        self.parent = parent
        self.file_path = None
        self.output_path = None
        self.df = pd.DataFrame()
        self.doc = Document()
        self.count_contrat_meditate = 0
        self.count_contrat_therapy = 0
        self.count_aditivo = 0
    
    def create_center_clinic_tab(self, center_clinic):
        layout_center_clinic = QVBoxLayout()

        # Layout horizontal para o QLabel e o botão "Limpar Status"
        status_layout = QHBoxLayout()

        # QLabel para exibir o status do arquivo
        self.label_status = QLabel("Nenhum arquivo carregado.")
        status_layout.addWidget(self.label_status)

        # Adiciona um espaço expansível para empurrar o botão para a direita
        status_layout.addStretch()

        # Botão para limpar o status
        btn_clear_status = QPushButton("Limpar Status")
        btn_clear_status.setFixedSize(150, 35)
        btn_clear_status.clicked.connect(self.clear_status)  # Conecta à função clear_status
        status_layout.addWidget(btn_clear_status)

        # Adiciona o layout horizontal ao layout vertical principal
        layout_center_clinic.addLayout(status_layout)

        # Adiciona os checkboxes para as opções
        checkbox_layout = QHBoxLayout()
        self.checkbox_group = QButtonGroup()
        self.checkbox_group.setExclusive(True)  # Apenas um pode ser selecionado

        self.checkbox_aditivo = QCheckBox("Aditivo")
        self.checkbox_aditivo.setChecked(True)  # Selecionado por padrão
        self.checkbox_group.addButton(self.checkbox_aditivo)
        checkbox_layout.addWidget(self.checkbox_aditivo)

        self.checkbox_contrato_medico = QCheckBox("Contrato Médico")
        self.checkbox_group.addButton(self.checkbox_contrato_medico)
        checkbox_layout.addWidget(self.checkbox_contrato_medico)

        self.checkbox_contratoterapia = QCheckBox("Contrato Terapia")
        self.checkbox_group.addButton(self.checkbox_contratoterapia)
        checkbox_layout.addWidget(self.checkbox_contratoterapia)

        layout_center_clinic.addLayout(checkbox_layout)

        # Separador horizontal
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        layout_center_clinic.addWidget(separator)

        # Layout horizontal para os botões
        button_layout = QHBoxLayout()

        # Botão para pesquisar o arquivo principal
        btn_search = QPushButton("Pesquisar")
        btn_search.setFixedSize(150, 35)
        btn_search.clicked.connect(self.search_file)  # Função de pesquisa (a ser implementada)
        button_layout.addWidget(btn_search)

        # Adiciona um espaço expansível entre os botões
        button_layout.addStretch()

        # Botão para processar e salvar o arquivo
        btn_process_save = QPushButton("Salvar")
        btn_process_save.setFixedSize(150, 35)
        btn_process_save.clicked.connect(self.process_and_save)
        button_layout.addWidget(btn_process_save)

        # Adiciona o layout horizontal ao layout vertical principal
        layout_center_clinic.addLayout(button_layout)

        # Configurando o layout na aba
        center_clinic.setLayout(layout_center_clinic)
    
    # Função para pesquisar o arquivo principal
    def search_file(self):
        # Cria uma instância da janela de pesquisa
        search_window = SearchWindow(self.parent)
        search_window.exec_()  # Abre a janela de pesquisa como modal

        self.df = search_window.df_search  # Atribui o DataFrame ao self.df

    def process_and_save(self):
        # Abre um diálogo para o usuário escolher a pasta onde os arquivos serão salvos
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        folder_path = QFileDialog.getExistingDirectory(
            self.parent,
            "Selecione a pasta para salvar os documentos",
            options=options
        )

        # Verifica se o usuário selecionou uma pasta
        if folder_path:
            self.output_path = folder_path  # Armazena o caminho da pasta em self.output_path

            # Verifica qual checkbox está selecionado e executa a função correspondente
            if self.checkbox_aditivo.isChecked():
                path_document = r'./ARQUIVOS/ADENDO NDI SP NDI RP.docx'
                self.count_aditivo = self.create_contract(path_document, name_arq='ADENDO NDI SP NDI', count_contract=self.count_aditivo)

            if self.checkbox_contrato_medico.isChecked():
                path_document = r'./ARQUIVOS/CONTRATO AMBULATORIAL NDI RP.docx'
                self.count_contrat_meditate = self.create_contract(path_document, name_arq='CONTRATO MÉDICO NDI', count_contract=self.count_contrat_meditate)

            if self.checkbox_contratoterapia.isChecked():
                path_document = r'./ARQUIVOS/CONTRATO EQUIPE MULTIDISCIPLINAR NDI RP.docx'
                self.count_contrat_therapy = self.create_contract(path_document, name_arq='CONTRATO TERAPIA NDI', count_contract=self.count_contrat_therapy)

            QMessageBox.information(self.parent, "Informação", f"Documentos serão salvos na pasta: {self.output_path}")
        else:
            QMessageBox.warning(self.parent, "Aviso", "Nenhuma pasta foi selecionada para salvar os documentos.")
        
    
    # Função para limpar o status
    def clear_status(self):
        #self.df = pd.DataFrame()
        self.label_status.setText("Nenhum arquivo carregado.")

    def create_contract(self, path_document, name_arq, count_contract):
        # Implementar a função de contrato médico
        protocol = self.df['CD_PROTOCOLO'].unique()
        #path_document = r'./ARQUIVOS/CONTRATO AMBULATORIAL NDI RP.docx'
        #doc = Document(path_document)
        df_especialidade = pd.read_csv(r'./ARQUIVOS/ESPECIALIDADE.csv', sep=';', encoding='latin1')
        # Salvando o arquivo
        for protocolo in protocol:
            # Criar o nome do arquivo baseado no protocolo
            doc = Document(path_document)
            count_contract += 1
            df_protocol = self.df[self.df['CD_PROTOCOLO'] == protocolo].copy()  # Filtra o DataFrame para o protocolo atual
            name_string = df_protocol['NM_RAZAO_NOME'].iloc[0]  # Obtém o nome do DataFrame
            date_doc = df_protocol['DT_INI_VIGENCIA'].iloc[0]  # Obtém a data do DataFrame

            name_file = f'{count_contract} - {name_arq}_{protocolo}.docx'

            # inserindo o protocolo no documento 
            self.replace_in_headers(doc=doc, old_word="XXX_XXX", new_word=str(protocolo))  # Substitui o texto no cabeçalho
            # Atualiza o texto no documento
            self.replace_text(doc, name_string, date_doc)  # Substitui o texto no documento
            print(name_string, protocolo)

            df_protocol = pd.merge(df_protocol, df_especialidade, how='left', left_on='CD_ESPECIALIDADE', right_on='COD_ESPECIALIDADE')
            df_protocol.CD_ORDEM_LOCAL = df_protocol.CD_ORDEM_LOCAL.astype(str).str.replace('.0', '', regex=False)
            df_protocol.CD_ORDEM_LOCAL = df_protocol.CD_ORDEM_LOCAL.astype(int)
            df = df_protocol[['ESPECIALIDADE', 'VL_HORA_PROPOSTO', 'NM_FANTASIA','CIDADE_UF']].copy()
            # df.DS_ESPECIALIDADE = df.DS_ESPECIALIDADE.str.capitalize()
            # df['CIDADE'] = df.CIDADE_UF.str.split('/').str[0].str.capitalize()
            # df['UF'] = df.CIDADE_UF.str.split('/').str[1].str.upper()
            # df.CIDADE_UF = df.CIDADE + '/' + df.UF
            # df.drop(columns=['CIDADE', 'UF'], inplace=True)
            df.sort_values(by=[ 'VL_HORA_PROPOSTO', 'CIDADE_UF'], ascending=[True, True], inplace=True)
            df.drop_duplicates(inplace=True)
            # df com após a virgula de duas unidades
            df.VL_HORA_PROPOSTO = df.VL_HORA_PROPOSTO.apply(lambda x: f"R$ {x:,.2f}".replace('.', ','))
            df.rename(
            columns={
                'ESPECIALIDADE': 'ESPECIALIDADE',
                'VL_HORA_PROPOSTO': 'VALOR HORA', 
                'NM_FANTASIA': 'LOCAL', 
                'CIDADE_UF': 'CIDADE/UF'
                }, 
            inplace=True
            )
            print(df.head(1), '\n')
            self.replace_table_content(doc=doc, table_index=2, df=df)  # Substitui o conteúdo da tabela

            # Definindo o caminho completo para salvar o arquivo
            path_save = os.path.join(self.output_path, name_file)

            # Salva o documento com o nome do protocolo
            doc.save(path_save)

        return count_contract


    # função para retornar a data de today em formato (dia, mês, ano)
    def date_today(self, value, date_doc):
        # Obtém a data de today
        today = datetime.now()
        string_date = None
        if value == 0:
            date_complete = today.strftime('%d/%m/%Y')
            date_complete = date_complete.replace('/', '_')
            string_date = date_complete
        
        if value == 1:
            date_doc = datetime.strptime(date_doc, '%Y-%m-%d %H:%M:%S')
            string_date = date_doc.strftime('%d/%m/%Y')
            day = date_doc.day
            month = date_doc.month
            dict_month = {
                1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
                5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
                9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
            }
            year = date_doc.year
            month = dict_month.get(month, 'Mês inválido')
            string_date = f'{day} de {month} de {year}'

        return string_date
    
    # Função para corrigir os nomes dos arquivos
    def replace_text(self, doc, name_string, date_doc):
           # Dicionários de substituição
        dict_replace = {
            'XX de XXX de 20XX': self.date_today(1, date_doc),
            '@NOME@': name_string,
        }
        # Iterando pelos parágrafos
        for pag, par in enumerate(doc.paragraphs):
            # Substitui os textos do dicionário dict_replace
            for key, value in dict_replace.items():
                if key in par.text:
                    for run in par.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

    # Função para adicionar bordas à tabela
    def add_table_borders(self, table):
        tbl = table._element
        tblPr = tbl.xpath("w:tblPr")[0]
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")  # Tipo de borda
            border.set(qn("w:sz"), "4")       # Espessura da borda
            border.set(qn("w:space"), "0")    # Espaçamento
            border.set(qn("w:color"), "000000")  # Cor da borda (preto)
            tblBorders.append(border)
        tblPr.append(tblBorders)

    # Função para substituir o conteúdo de uma tabela específica
    def replace_table_content(self, doc, table_index, df):
        table = doc.tables[table_index]

        # Substitui o cabeçalho da tabela
        header_cells = table.rows[0].cells
        for col_index, col_name in enumerate(df.columns):
            cell = header_cells[col_index]
            cell.text = col_name
            # Formatação do cabeçalho
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centralizado
            run = paragraph.runs[0]
            run.bold = True  # Negrito
            run.font.size = Pt(10)  # Tamanho da fonte

        # Remove as linhas existentes (exceto o cabeçalho)
        for row in table.rows[1:]:
            tr = row._element
            tr.getparent().remove(tr)

        # Adiciona as novas linhas com os dados do DataFrame
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for col_index, value in enumerate(row):
                cell = row_cells[col_index]
                cell.text = str(value)
                # Formatação das células
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centralizado
                run = paragraph.runs[0]
                run.font.size = Pt(10)  # Tamanho da fonte

        # Adiciona bordas à tabela
        self.add_table_borders(table)
    
    # Função para substituir palavras nos cabeçalhos sem mudar o conteúdo
    def replace_in_headers(self, doc, old_word, new_word):
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    if old_word in run.text: # Verifica se o texto antigo está na run
                        run.text = run.text.replace(old_word, new_word) # Substitui o texto na run


class SearchWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Pesquisar Dados")
        self.setFixedSize(600, 400)  # Tamanho fixo da janela
        self.init_ui()
        self.df_search = pd.DataFrame()  # DataFrame para armazenar os dados pesquisados
        self.parent = parent  # Armazena a referência ao widget pai

    def init_ui(self):
        # Cria o layout principal
        main_layout = QVBoxLayout()

        # Linha para entrada de texto e botão de pesquisa
        search_layout = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Digite os Protocolos...")
        search_layout.addWidget(self.search_input)

        btn_search = QPushButton("Pesquisar")
        btn_search.setFixedSize(100, 30)
        btn_search.clicked.connect(self.perform_search)  # Conecta à função de pesquisa
        search_layout.addWidget(btn_search)

        main_layout.addLayout(search_layout)

        # Separador
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setFrameShadow(QFrame.Sunken)
        main_layout.addWidget(separator1)

        # Tabela para exibir os resultados
        self.table = QTableWidget()
        self.table.setColumnCount(14)  # Define 14 colunas
        self.table.setHorizontalHeaderLabels([
            "CD_PROTOCOLO", "NM_RAZAO_NOME", "DT_INI_VIGENCIA",
            "CD_ESPECIALIDADE", "DS_ESPECIALIDADE","VL_HORA_PROPOSTO", "CD_LOCAL",
            "CD_ORDEM_LOCAL", "NU_ORDEM_LOCA", "FL_URGENCIA",
            "FL_ELETIVA", "CD_LOCAL_ATENDIMENTO", "NM_FANTASIA", "CD_UF"
        ])
        #self.table.setRowCount()  # Exibe inicialmente 5 linhas
        main_layout.addWidget(self.table)

        # Separador
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.HLine)
        separator2.setFrameShadow(QFrame.Sunken)
        main_layout.addWidget(separator2)

        # Botão para armazenar a informação (centralizado)
        btn_store_layout = QHBoxLayout()
        btn_store = QPushButton("Armazenar Informação")
        btn_store.setFixedSize(200, 40)
        btn_store.clicked.connect(self.store_information)  # Conecta à função de armazenamento
        btn_store_layout.addStretch()  # Adiciona espaço antes do botão
        btn_store_layout.addWidget(btn_store)
        btn_store_layout.addStretch()  # Adiciona espaço depois do botão
        main_layout.addLayout(btn_store_layout)

        # Adiciona o layout principal a um widget para o QScrollArea
        container_widget = QWidget()
        container_widget.setLayout(main_layout)

        # Cria uma área de rolagem
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(container_widget)

        # Define o layout da janela
        window_layout = QVBoxLayout()
        window_layout.addWidget(scroll_area)
        self.setLayout(window_layout)

    def perform_search(self):
        # Obtém o termo de pesquisa e o caminho do driver JDBC
        search_term = self.search_input.text()
        path_drive = r'./ARQUIVOS/Oracle_Jdbc/ojdbc8.jar'

        if search_term:
            try:
                # Instancia a classe JdbcPermission
                jdbc_permission = JdbcPermission(path_drive)

                # Usa o método fetch_data para buscar os dados
                df, protocol = jdbc_permission.fetch_data(search_term)

                cd_protocol = df.CD_PROTOCOLO.unique().tolist()

                protocol = protocol.split(', ')

                protocol = [int(p) for p in protocol]  # Converte os protocolos para inteiros

                protocol_diff = list(set(protocol) - set(cd_protocol))

                print(protocol, cd_protocol, protocol_diff)

                # Atualiza o self.df_search com os dados encontrados
                self.df_search = df
                # Define o número de linhas e colunas da tabela com base no DataFrame
                self.table.setRowCount(len(df))
                self.table.setColumnCount(len(df.columns))
                self.table.setHorizontalHeaderLabels(df.columns)

                # Preenche a tabela com os dados do DataFrame
                for row_idx, row_data in df.iterrows():
                    for col_idx, cell_data in enumerate(row_data):
                        self.table.setItem(row_idx, col_idx, QTableWidgetItem(str(cell_data)))

                if self.df_search.empty:
                    QMessageBox.warning(self, 'Aviso', f' 1 - Protocolo(s): ( {protocol} ) inelegível para automatização de documentos. \n\nSolicita-se verificação com a coordenação.')

                if len(protocol) != len(cd_protocol):
                    protocol_diff = ', '.join([str(p) for p in protocol_diff])
                    QMessageBox.warning(self, 'Aviso', f'2 - Protocolo(s): ( {protocol_diff} ) inelegível para automatização de documentos. \n\nSolicita-se verificação com a coordenação.')


            except Exception as erro:
                QMessageBox.critical(self, "Erro", f"Erro ao buscar dados:\n{str(erro)}")
        else:
            QMessageBox.warning(self, "Aviso", "Digite um termo para pesquisar!")

    def store_information(self):
        # Verifica se há dados no DataFrame
        if self.df_search.empty:
            QMessageBox.warning(self, "Aviso", "Nenhuma informação foi encontrada para armazenar!")
            return None

        # Exibe uma mensagem de sucesso e retorna o DataFrame
        QMessageBox.information(self, "Informação", "Informação armazenada com sucesso!")
        return self.df_search